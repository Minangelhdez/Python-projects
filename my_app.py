from docx import Document
from docx.shared import Inches
import pyttsx3 

def speak(text):
    pyttsx3.speak(text)

document = Document()


# Profile picture
document.add_picture('me.jpg', width=Inches(2.0))

# Name phone number and email details
name = input('What is your full name? ')
speak('Hello ' + name + ' how are you today ? ')
phone_number = input('Thanks, what is your phone number? ')
email = input('what is email address? ')


document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email)

# About me
document.add_heading('About me')
about_me= input('Tell me about yourself. ')
document.add_paragraph(about_me)

# Work experience 
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter company ')
from_date = input('From Date ')
to_date = input('To Date ')

p.add_run( company + '  ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

experience_details = input(

    'Decribe your experience at  ' + company)
p.add_run (experience_details)

#more experiences
while True: 
    has_more_experiences= input(
        'Do you have more experiences ?(Yes or No)  ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter company ')
        from_date = input('From Date ')
        to_date = input('To Date ')

        p.add_run( company + '  ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        experience_details = input(

        'Decribe your experience at ' + company + ' ')
        p.add_run (experience_details)
    else:
        break    

#Skills
document.add_heading(' Skills ')
skills = input('What skills do you have? ')
p = document.add_paragraph(skills)
p.style = 'List Bullet'

#more skills
while True:
    has_more_skills = input('Nice, Do you have any other skills? (Yes or No)  ')
    if has_more_skills.lower() == 'yes':
        skills = input('What other skill do you have? ')
        p = document.add_paragraph(skills)
        p.style = 'List Bullet'

    else :
        break




print('   ')
print('-----FILE HAVE BEEN UPDATED-----!')
document.save('cv.docx')